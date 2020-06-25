#!/usr/bin/env python
# Copyright (c) 2018 The Regents of the University of Michigan
# All rights reserved.
# This software is licensed under the BSD 3-Clause License.

"""This is a helper script to generate a .docx file containing all templates for easy review."""
import os
from operator import attrgetter
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import git

import signac

import extract_templates as ext

DOC_DIR = os.path.join(
    os.path.dirname(__file__), 'compiled_scripts')


def add_param_header(document, job):
    h = ['{}={}'.format(key, val) for key, val in job.sp.parameters.items()]
    h = 'Parameters: ' + ', '.join(h)
    document.add_heading(h, level=2)


def process_job(document, job):
    add_param_header(document, job)

    from collections import OrderedDict
    name_map = OrderedDict([
        ('script_serial_op.sh',  'Serial operation'),
        ('script_parallel_op.sh',  'Generic parallel operation'),
        ('script_mpi_op.sh',  'MPI operation'),
        ('script_omp_op.sh',  'OpenMP operation'),
        ('script_hybrid_op.sh',  'MPI-OpenMP Hybrid operation'),
        ('script_gpu_op.sh',  'GPU operation'),
        ('script_mpi_gpu_op.sh',  'MPI-GPU operation'),
        ('script_group1.sh',  'Operation group'),
    ])

    if job.sp.parameters.get('bundle', False):
        h = 'Bundled MPI and OpenMP jobs'
        for fn in os.listdir(job.workspace()):
            if fn == 'signac_statepoint.json':
                continue
            document.add_heading(h, level=3)
            with open(job.fn(fn)) as fr:
                document.add_paragraph(fr.read(), style='Code')
    else:
        for fn, h in name_map.items():
            if job.isfile(fn):
                document.add_heading('Operation: {}'.format(h), level=3)
                with open(job.fn(fn)) as fr:
                    document.add_paragraph(fr.read(), style='Code')

    document.add_page_break()


def main():
    repo = git.Repo(search_parent_directories=True)
    commit = repo.head.commit

    with open('define_template_test_project.py', 'r') as file:
        PROJECT_DEFINITION = file.read()

    if not os.path.exists(DOC_DIR):
        os.makedirs(DOC_DIR)
    try:
        project = signac.get_project(ext.PROJECT_DIR)
    except LookupError:
        print('The templates project could not be found. Try running '
              './extract_templates.py before this script.')
        raise

    sort_keys = list(sorted(project.detect_schema()))
    sort_attrs = ['sp.' + '.'.join(key) for key in sort_keys]

    def sort_key(job):
        values = []
        for attr in sort_attrs:
            try:
                values.append(attrgetter(attr)(job))
            except AttributeError:
                pass
        return [0 if v is None else v for v in values]

    environments = project.detect_schema()['environment'][str]
    for env in sorted(environments):
        env_name = env.split('.')[-1]
        document = docx.Document()

        # Add code style
        style = document.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Monaco'
        style.font.size = Pt(8)
        style = document.styles.add_style('CodeChar', WD_STYLE_TYPE.CHARACTER)
        style.font.name = 'Monaco'
        style.font.size = Pt(8)

        document.add_heading(env_name, level=0)
        p = document.add_paragraph("Output at commit ")
        p.add_run('{}'.format(commit), style='CodeChar')
        document.add_heading("FlowProject Definition", level=1)
        p = document.add_paragraph(PROJECT_DEFINITION, style='Code')
        document.add_page_break()

        document.add_heading("Operations without bundling", level=1)
        query = {'environment': env, 'parameters.bundle': {'$exists': False}}
        for job in sorted(project.find_jobs(query), key=sort_key):
            process_job(document, job)

        document.add_heading("Operations with bundling", level=1)
        query = {'environment': env, 'parameters.bundle': {'$exists': True}}
        for job in sorted(project.find_jobs(query), key=sort_key):
            process_job(document, job)

        fn = os.path.join(DOC_DIR, "{env}.docx".format(env=env_name))
        document.save(fn)
        print("Generated document '{}'.".format(fn))


if __name__ == "__main__":
    main()
